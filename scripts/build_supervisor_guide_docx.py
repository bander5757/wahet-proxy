from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/دليل-المشرف-واحة-الخيمة.docx")


def set_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    for run in paragraph.runs:
        run.font.name = "Tahoma"
        run._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")


def add_paragraph(doc, text="", style=None, bold=False, color=None, size=None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        if size:
            run.font.size = Pt(size)
    set_rtl(p)
    return p


def add_heading(doc, text, level=1):
    p = add_paragraph(doc, text, style=f"Heading {level}")
    for run in p.runs:
        run.font.name = "Tahoma"
        run._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
        run.font.color.rgb = RGBColor(31, 111, 104) if level == 1 else RGBColor(184, 121, 47)
    return p


def add_bullets(doc, items):
    for item in items:
        p = add_paragraph(doc, item, style="List Bullet")
        set_rtl(p)


def add_steps(doc, items):
    for item in items:
        p = add_paragraph(doc, item, style="List Number")
        set_rtl(p)


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "FFF8EB")
    cell._tc.get_or_add_tcPr().append(shading)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(title + ": ")
    r.bold = True
    r.font.name = "Tahoma"
    r._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
    r2 = p.add_run(body)
    r2.font.name = "Tahoma"
    r2._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
    set_rtl(p)
    add_paragraph(doc, "")


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Tahoma"
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.name = "Tahoma"
    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 2"].font.name = "Tahoma"
    styles["Heading 2"].font.size = Pt(14)

    add_paragraph(doc, "دليل المشرف - برنامج واحة الخيمة", bold=True, color=(31, 111, 104), size=20)
    add_paragraph(doc, "شرح مبسط لعمر كمدخل بيانات ومتابع للعمل اليومي من الجوال أو الكمبيوتر.", color=(86, 98, 114), size=11)
    add_note(doc, "الهدف", "تسجيل البيانات بوضوح، متابعة العملاء، وتنبيه بندر وابو فايز بأي شيء يحتاج قرار.")
    add_note(doc, "أهم قاعدة", "لا تسجل أي مبلغ أو حركة بدون ملاحظة واضحة توضح السبب والمستلم أو العميل.")

    add_heading(doc, "1. الدخول للبرنامج")
    add_steps(doc, [
        "افتح رابط البرنامج.",
        "اضغط زر دخول.",
        "اختر حساب المشرف.",
        "اكتب كلمة المرور.",
        "إذا كان الجهاز خاص بك، فعّل حفظ الدخول حتى لا تحتاج تسجيل الدخول كل مرة.",
    ])

    add_heading(doc, "2. الرئيسية")
    add_paragraph(doc, "هذه الصفحة تعطيك نظرة سريعة على وضع المؤسسة.")
    add_bullets(doc, [
        "راجع العروض والفواتير القادمة.",
        "راجع المبالغ المسددة والمتبقية.",
        "انتبه للتركيبات القريبة.",
        "راجع التنبيهات المهمة للموظفين والسيارات.",
        "إذا ظهر شيء يحتاج متابعة، افتح التبويب المناسب وسجل الإجراء.",
    ])

    add_heading(doc, "3. العملاء")
    add_paragraph(doc, "هذا التبويب يعرض العملاء وعروض الأسعار والفواتير القادمة من دفترة.")
    add_bullets(doc, [
        "ابدأ دائمًا بعروض الشهر الحالي.",
        "العرض غير المؤكد يحتاج اتصال أو متابعة مع العميل.",
        "إذا وافق العميل، غيّر حالة العرض إلى مؤكد.",
        "إذا العميل اعتذر أو استغنى، غيّر الحالة حتى لا يبقى العرض مفتوحًا.",
        "افتح تفاصيل العرض لمراجعة رقم العميل، الموقع، المدة، وتاريخ التركيب إذا كانت موجودة.",
    ])
    add_note(doc, "ملاحظة", "إعدادات ربط دفترة خاصة ببندر فقط، لا تحتاج تعديلها.")

    add_heading(doc, "4. المحاسبة")
    add_paragraph(doc, "هذا التبويب لتسجيل الحركات المالية اليومية.")
    add_heading(doc, "الحركات التي تسجلها", 2)
    add_bullets(doc, ["مصروف.", "عهدة.", "إيراد.", "سلفة.", "دين أو مبلغ متبقٍ."])
    add_heading(doc, "عند التسجيل", 2)
    add_bullets(doc, [
        "اكتب المبلغ بدقة.",
        "اختر الحساب: الرسمي أو الفرعي.",
        "اكتب الشخص أو العميل المرتبط بالحركة.",
        "اكتب ملاحظة واضحة: سبب التحويل، المستلم، ورقم الفاتورة إن وجد.",
        "أرفق صورة الإيصال إذا كانت متوفرة.",
    ])

    add_heading(doc, "5. المحاسب الذكي")
    add_paragraph(doc, "هذا الجزء يساعد في مراجعة كشف البنك ومطابقة الحركات.")
    add_bullets(doc, [
        "الأفضل رفع كشف البنك بصيغة CSV أو ملف نصي مرتب.",
        "PDF يمكن حفظه للتوثيق، لكن التحليل الدقيق منه يحتاج تطوير إضافي لاستخراج الجدول.",
        "التحليل يعطي مطابقة مبدئية وتقرير أسبوعي وتقدير أولي للضريبة.",
    ])
    add_note(doc, "تنبيه", "لا تعتمد على المحاسب الذكي وحده للرفع الضريبي النهائي إلا بعد مراجعة بندر أو المحاسب الزائر.")

    add_heading(doc, "6. التركيبات")
    add_bullets(doc, [
        "أي فاتورة من دفترة تظهر ضمن التركيبات.",
        "أي عرض سعر تم وضعه كمؤكد يظهر ضمن التركيبات.",
        "راجع تاريخ التركيب، موقع العميل، رقم الجوال، والمدة.",
        "إذا كانت البيانات ناقصة، راجع تفاصيل العرض أو حدّث ملاحظات دفترة.",
    ])

    add_heading(doc, "7. الموظفين والسيارات")
    add_bullets(doc, [
        "انتهاء الإقامات.",
        "انتهاء رخص العمل.",
        "صيانة السيارات.",
        "تغيير الزيت.",
        "أي تنبيه عام مثل موعد مراجعة أو تجديد أو مهمة إدارية.",
    ])
    add_paragraph(doc, "عند إضافة تنبيه، اكتب اسم التنبيه، تاريخ التنبيه، وملاحظة واضحة.")

    add_heading(doc, "8. المنافسات")
    add_bullets(doc, [
        "ركز على خيام أوروبية، خيم، خيمة أوروبية، وتجهيز فعاليات تحتاج خيام.",
        "افتح رابط المصدر وراجع التفاصيل.",
        "سجل القرار: مناسبة، غير مناسبة، أو تحتاج مراجعة.",
        "إذا كانت مناسبة، بلغ بندر للمراجعة قبل التقديم.",
    ])

    add_heading(doc, "9. التقارير")
    add_bullets(doc, [
        "تقرير أسبوعي.",
        "تقرير شهري.",
        "ملخص المصاريف.",
        "ملخص الإيرادات.",
        "المبالغ المتبقية.",
    ])
    add_paragraph(doc, "إذا ظهر رقم غير واضح، راجع مصدر الحركة قبل التعديل.")

    add_heading(doc, "10. زر حركة سريع")
    add_paragraph(doc, "زر + حركة مخصص للتسجيل السريع من الجوال.")
    add_bullets(doc, [
        "استخدمه عند تسجيل مصروف أو حوالة أو عهدة أو إيراد.",
        "اكتب ملاحظة قصيرة وواضحة.",
        "أرفق الإيصال إذا توفر.",
    ])

    add_heading(doc, "طريقة العمل اليومية للمشرف")
    add_steps(doc, [
        "افتح الرئيسية وشوف التنبيهات والمهم اليوم.",
        "افتح العملاء وراجع عروض الشهر الحالي.",
        "تابع العروض غير المؤكدة مع العملاء.",
        "سجل أي مصروف أو عهدة أو حوالة فورًا من زر حركة.",
        "راجع التركيبات القادمة وتأكد أن بياناتها واضحة.",
        "بلغ بندر بأي رقم ناقص، عرض غير واضح، أو منافسة مناسبة.",
    ])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
